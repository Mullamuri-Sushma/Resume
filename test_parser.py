from docx import Document
import zipfile
import os

# create sample docx
doc = Document()
doc.add_heading('John Doe', level=1)
doc.add_paragraph('Email: john.doe@example.com')
doc.add_paragraph('Phone: +1 (555) 123-4567')
doc.add_paragraph('\nSkills:\nPython, pandas, SQL, Docker')
doc.add_paragraph('\nExperience:\nSenior Software Engineer at Acme Corp. Worked on data pipelines and ML features.')
doc.add_paragraph('\nEducation:\nB.Sc Computer Science, University of Somewhere, 2016-2020')
doc.add_paragraph('\nLinkedIn: https://www.linkedin.com/in/johndoe')
doc.add_paragraph('GitHub: https://github.com/johndoe')

sample_docx = 'sample_resume.docx'
doc.save(sample_docx)

# create a zip containing the docx
zip_name = 'sample_resumes.zip'
with zipfile.ZipFile(zip_name, 'w') as z:
    z.write(sample_docx)

print('Created', sample_docx, 'and', zip_name)

# now run parser by importing website (module) and using its helpers
import importlib
import website

text = website.read_docx(sample_docx)
parsed = website.process_resume(text)
print('Parsed result:')
print(parsed)

# cleanup
os.remove(sample_docx)
os.remove(zip_name)
